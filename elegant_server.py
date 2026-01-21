#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from flask import Flask, request, jsonify, send_from_directory, make_response
from flask_cors import CORS
import html_report
import os
import pdfplumber
from PIL import Image
import pytesseract
import numpy as np
import requests
import io
import base64
import time
import uuid
import tempfile
import functools

# 加载环境变量
from dotenv import load_dotenv
load_dotenv()

# 导入大模型API代理层
from llm_proxy import llm_proxy

# 文件存储管理
temp_files = {}

# 解决paddleocr中的langchain导入问题
# 在新版本的langchain中，许多模块已经被移动到langchain-community或独立的包中
try:
    import sys
    
    # 设置langchain.docstore.document别名
    from langchain_community.docstore.document import Document
    sys.modules['langchain.docstore.document'] = sys.modules['langchain_community.docstore.document']
    print("已设置langchain.docstore.document别名")
    
    # 设置langchain.text_splitter别名，使用独立的langchain-text-splitters包
    try:
        from langchain_text_splitters import RecursiveCharacterTextSplitter
        sys.modules['langchain.text_splitter'] = sys.modules['langchain_text_splitters']
        print("已设置langchain.text_splitter别名")
    except ImportError:
        print("langchain-text-splitters包未安装，跳过设置text_splitter别名")
    
    # 暂时注释掉paddleocr导入，避免模型源检查
    # try:
    #     from paddleocr import PaddleOCRVL
    #     # 只导入，不创建实例，避免初始化PDX
    #     print("已导入PaddleOCRVL")
    # except Exception as e:
    #     print(f"导入PaddleOCRVL失败: {e}")
except Exception as e:
    print(f"设置langchain别名失败: {e}")
    import traceback
    traceback.print_exc()

app = Flask(__name__)
CORS(app)  # 添加CORS支持

# 设置静态文件目录
app.static_folder = '.'

# 频率限流配置
REQUEST_LIMIT = 3  # 每分钟最大请求数
request_history = {}

# 频率限流装饰器
def rate_limit(func):
    @functools.wraps(func)
    def wrapper(*args, **kwargs):
        # 获取客户端IP
        client_ip = request.remote_addr
        current_time = int(time.time() / 60)  # 按分钟计数
        
        # 初始化请求历史
        if client_ip not in request_history:
            request_history[client_ip] = {}
        
        # 清理过期记录
        expired_times = [t for t in request_history[client_ip] if t < current_time - 1]
        for t in expired_times:
            del request_history[client_ip][t]
        
        # 检查请求频率
        if current_time in request_history[client_ip]:
            if request_history[client_ip][current_time] >= REQUEST_LIMIT:
                return jsonify({'error': '请求过于频繁，请稍后重试'}), 429
        else:
            request_history[client_ip][current_time] = 0
        
        # 增加请求计数
        request_history[client_ip][current_time] += 1
        
        return func(*args, **kwargs)
    return wrapper

def ocr_olmocr(image):
    """
    使用OCR技术进行文字识别，优先使用pytesseract（CPU友好）
    """
    # 保存图像到临时文件
    import tempfile
    import os
    
    with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp:
        image.save(tmp, format="PNG")
        tmp_path = tmp.name
    
    try:
        # 优先使用pytesseract（CPU友好的OCR技术）
        try:
            result = pytesseract.image_to_string(image, lang='chi_sim+eng')
            if result.strip():
                return result.strip()
            else:
                return "（OCR识别成功，但未提取到文本）"
        except ImportError:
            print("pytesseract库未安装，无法使用OCR功能")
            return "（OCR识别失败，请先安装tesseract和pytesseract库）"
        except Exception as e:
            print(f"pytesseract识别失败: {e}")
            
            # 尝试使用EasyOCR作为备用方案（也是CPU友好的）
            try:
                import easyocr
                
                # 创建EasyOCR实例（指定语言：中文和英文）
                # 指定临时目录存储模型文件，解决权限问题
                reader = easyocr.Reader(
                    ['ch_sim', 'en'], 
                    gpu=False,  # 禁用GPU，使用CPU
                    download_enabled=False,  # 禁用自动下载
                    user_network_directory='/tmp/EasyOCR'  # 指定可写的临时目录
                )
                
                # 调用readtext方法，传递图像路径
                result = reader.readtext(tmp_path)
                
                # 处理返回结果
                text = ""
                for detection in result:
                    if len(detection) >= 2:
                        text += detection[1] + "\n"
                
                if text.strip():
                    return text.strip()
                else:
                    return "（OCR识别成功，但未提取到文本）"
            except ImportError:
                print("easyocr库未安装，无法使用备用OCR功能")
                return "（OCR识别失败，请先安装tesseract和pytesseract库或easyocr库）"
            except Exception as e:
                print(f"easyocr识别失败: {e}")
                return "（OCR识别失败，请安装并配置好OCR环境）"
    finally:
        # 删除临时文件
        os.remove(tmp_path)

def extract_file_content(file):
    """
    提取不同类型文件的内容
    :param file: 上传的文件对象
    :return: 文件内容文本
    """
    if not file:
        return None
    
    import os
    
    try:
        file_extension = os.path.splitext(file.filename)[1].lower()
        
        if file_extension == '.pdf':
            import pdfplumber
            text = ""
            with pdfplumber.open(file) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text
                    else:
                        # 如果文本提取失败，尝试OCR
                        ocr_result = ocr_olmocr(page.to_image().original)
                        # 无论OCR结果如何，都将其添加到文本中
                        text += ocr_result
            # 返回提取的文本，即使为空
            return text
        elif file_extension in ['.png', '.jpg', '.jpeg', '.bmp', '.gif']:
            from PIL import Image
            image = Image.open(file)
            ocr_result = ocr_olmocr(image)
            # 无论OCR结果如何，都返回结果
            return ocr_result
        elif file_extension == '.docx':
            import docx
            doc = docx.Document(file)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text  # 返回提取的文本，即使为空
        elif file_extension == '.doc':
            import tempfile
            
            # 保存临时文件
            file_content = file.read()  # 读取文件内容
            with tempfile.NamedTemporaryFile(delete=False, suffix='.doc') as tmp:
                tmp.write(file_content)
                tmp_path = tmp.name
            
            # 重置文件指针到开头，以便后续可能的读取操作
            file.seek(0)
            
            try:
                import win32com.client as win32
                
                # 使用Word应用程序打开文件
                word = win32.Dispatch('Word.Application')
                word.Visible = False
                doc = word.Documents.Open(tmp_path)
                text = doc.Content.Text
                doc.Close()
                word.Quit()
                return text  # 返回提取的文本，即使为空
            except:
                try:
                    import subprocess
                    
                    # 使用LibreOffice转换.doc为.docx
                    output_path = tmp_path + 'x'
                    subprocess.run([
                        'soffice', '--headless', '--convert-to', 'docx', 
                        tmp_path, '--outdir', os.path.dirname(tmp_path)
                    ], check=True)
                    
                    # 读取转换后的docx文件
                    import docx
                    doc = docx.Document(output_path)
                    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                    
                    # 删除临时文件
                    os.remove(output_path)
                    return text  # 返回提取的文本，即使为空
                except Exception as e:
                    print(f"无法读取.doc文件: {e}")
                    return "（无法读取.doc文件，请尝试转换为.docx格式后再上传）"  # 返回错误信息
            finally:
                # 删除临时文件
                os.remove(tmp_path)
        # 处理文本文件
        try:
            content = file.read().decode('utf-8')
            return content  # 返回提取的文本，即使为空
        except UnicodeDecodeError:
            # 尝试使用其他编码
            try:
                file.seek(0)
                content = file.read().decode('gbk')
                return content  # 返回提取的文本，即使为空
            except Exception as e:
                print(f"读取文本文件时出错: {e}")
                return "（文本文件编码错误，无法读取内容）"  # 返回错误信息
    
    except Exception as e:
        print(f"读取文件时出错: {e}")
        return "（文件读取失败，请检查文件格式和内容）"  # 返回错误信息

def analyze_resume_with_AI(resume_text: str, jd_text: str, target_position: str) -> str:
    """
    使用AI分析简历内容
    :param resume_text: 简历文本内容
    :param jd_text: 职位JD文本内容
    :param target_position: 目标岗位
    :return: AI分析结果
    """
    if not llm_proxy:
        return "服务暂时不可用，请稍后重试"
    
    try:
        # 使用大模型API代理层进行分析
        result = llm_proxy.analyze_resume(resume_text, jd_text, target_position)
        return result
    except Exception as e:
        # 避免泄露密钥相关错误信息
        print(f"大模型API调用失败：{str(e)}")
        return "服务暂时不可用，请稍后重试"

@app.route('/')
def index():
    return send_from_directory('.', 'elegant_app.html')

@app.route('/analyze', methods=['POST'])
@rate_limit
def analyze():
    try:
        # 获取请求数据
        target_position = request.form.get('target_position')
        
        # 获取JD内容
        jd_text = request.form.get('jd_text', '')
        jd_file = request.files.get('jd_file')
        jd_content = jd_text
        
        if jd_file:
            jd_content = extract_file_content(jd_file)
        
        # 获取简历内容
        resume_text = request.form.get('resume_text', '')
        resume_file = request.files.get('resume_file')
        resume_content = resume_text
        
        if resume_file:
            resume_content = extract_file_content(resume_file)
        
        # 验证数据
        if not target_position:
            return jsonify({'error': '请选择目标岗位'}), 400
        if not jd_text and not jd_file:
            return jsonify({'error': '请输入职位JD内容或上传有效的JD文件'}), 400
        if not resume_text and not resume_file:
            return jsonify({'error': '请输入简历内容或上传有效的简历文件'}), 400
        # 允许OCR识别失败的情况通过验证，让AI分析来处理这些情况
        if not jd_content:
            jd_content = "（OCR识别失败，无法提取JD图像中的文字内容）"
        if not resume_content:
            resume_content = "（OCR识别失败，无法提取简历图像中的文字内容）"
        
        # 调用AI分析
        analysis_result = analyze_resume_with_AI(resume_content, jd_content, target_position)
        
        # 返回结果
        return jsonify({'analysis': analysis_result})
    except Exception as e:
        # 避免泄露详细错误信息
        print(f"分析请求处理失败：{str(e)}")
        return jsonify({'error': '服务暂时不可用，请稍后重试'}), 500

@app.route('/upload', methods=['POST'])
@rate_limit
def handle_file_upload():
    """
    文件上传接口
    返回临时文件ID，用于后续分析请求
    """
    try:
        file = request.files.get('file')
        if not file:
            return jsonify({'error': '请选择要上传的文件'}), 400
        
        # 生成唯一文件ID
        file_id = str(uuid.uuid4())
        
        # 提取文件内容
        file_content = extract_file_content(file)
        
        # 存储文件内容
        temp_files[file_id] = {
            'content': file_content,
            'filename': file.filename,
            'timestamp': time.time()
        }
        
        # 清理过期文件（超过1小时的文件）
        current_time = time.time()
        expired_files = [fid for fid, data in temp_files.items() if current_time - data['timestamp'] > 3600]
        for fid in expired_files:
            del temp_files[fid]
        
        return jsonify({'file_id': file_id})
    except Exception as e:
        print(f"文件上传失败：{str(e)}")
        return jsonify({'error': '文件上传失败，请稍后重试'}), 500

@app.route('/api/analysis/start', methods=['POST'])
@rate_limit
def start_analysis():
    """
    分析初始化接口
    接收JD文件ID、简历文件ID和目标岗位，启动分析流程
    """
    try:
        # 获取请求数据
        data = request.json
        if not data:
            return jsonify({'error': '请求数据不能为空'}), 400
        
        target_position = data.get('target_position')
        jd_file_id = data.get('jd_file_id')
        resume_file_id = data.get('resume_file_id')
        jd_text = data.get('jd_text', '')
        resume_text = data.get('resume_text', '')
        
        # 验证数据
        if not target_position:
            return jsonify({'error': '请选择目标岗位'}), 400
        
        # 获取JD内容
        jd_content = jd_text
        if jd_file_id:
            if jd_file_id not in temp_files:
                return jsonify({'error': 'JD文件ID无效或已过期'}), 400
            jd_content = temp_files[jd_file_id]['content']
        
        # 获取简历内容
        resume_content = resume_text
        if resume_file_id:
            if resume_file_id not in temp_files:
                return jsonify({'error': '简历文件ID无效或已过期'}), 400
            resume_content = temp_files[resume_file_id]['content']
        
        # 验证内容
        if not jd_content:
            return jsonify({'error': '请输入职位JD内容或上传有效的JD文件'}), 400
        if not resume_content:
            return jsonify({'error': '请输入简历内容或上传有效的简历文件'}), 400
        
        # 调用AI分析
        analysis_result = analyze_resume_with_AI(resume_content, jd_content, target_position)
        
        # 生成HTML报告
        html_content = html_report.markdown_to_html(analysis_result, target_position)
        
        return jsonify({'code': 200, 'data': {'report': html_content}})
    except Exception as e:
        # 避免泄露详细错误信息
        print(f"分析初始化失败：{str(e)}")
        return jsonify({'code': 500, 'msg': '服务暂时不可用，请稍后重试'})

@app.route('/analyze/html', methods=['POST'])
@rate_limit
def analyze_html():
    try:
        # 获取请求数据
        target_position = request.form.get('target_position')
        
        # 获取JD内容
        jd_text = request.form.get('jd_text', '')
        jd_file = request.files.get('jd_file')
        jd_content = jd_text
        
        if jd_file:
            jd_content = extract_file_content(jd_file)
        
        # 获取简历内容
        resume_text = request.form.get('resume_text', '')
        resume_file = request.files.get('resume_file')
        resume_content = resume_text
        
        if resume_file:
            resume_content = extract_file_content(resume_file)
        
        # 验证数据
        if not target_position:
            return jsonify({'error': '请选择目标岗位'}), 400
        if not jd_text and not jd_file:
            return jsonify({'error': '请输入职位JD内容或上传有效的JD文件'}), 400
        if not resume_text and not resume_file:
            return jsonify({'error': '请输入简历内容或上传有效的简历文件'}), 400
        # 允许OCR识别失败的情况通过验证，让AI分析来处理这些情况
        if not jd_content:
            jd_content = "（OCR识别失败，无法提取JD图像中的文字内容）"
        if not resume_content:
            resume_content = "（OCR识别失败，无法提取简历图像中的文字内容）"
        
        # 调用AI分析
        analysis_result = analyze_resume_with_AI(resume_content, jd_content, target_position)
        
        # 生成HTML报告
        html_content = html_report.markdown_to_html(analysis_result, target_position)
        
        # 创建响应
        response = make_response(html_content)
        response.headers['Content-Type'] = 'text/html; charset=utf-8'
        
        return response
    except Exception as e:
        # 避免泄露详细错误信息
        print(f"HTML报告生成失败：{str(e)}")
        return jsonify({'error': '服务暂时不可用，请稍后重试'}), 500

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 8888))
    print(f"\n=== AIPM 简历分析智能助手 ===")
    print(f"服务器正在运行...")
    print(f"\n请在浏览器中访问: http://localhost:{port}")
    print(f"\n按 Ctrl+C 停止服务器")
    print(f"================================\n")
    app.run(host='0.0.0.0', port=port, debug=False)
