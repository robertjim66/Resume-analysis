from openai import OpenAI
import pdfplumber
from PIL import Image
import pytesseract
import numpy as np
import requests
import io
import base64
from flask import Flask, render_template_string, request, redirect, url_for, flash, session
from paddleocr import PaddleOCRVL

# 创建Flask应用
app = Flask(__name__)
app.secret_key = 'supersecretkey'  # 用于session管理

# 目标岗位改为用户自定义输入

def ocr_olmocr(image):
    """
    使用OCR OlmOCR进行文字识别
    """
    try:
        # 将图像转换为base64
        buffer = io.BytesIO()
        image.save(buffer, format="PNG")
        img_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')
        
        # 使用PaddleOCRVL进行OCR识别
        pipeline = PaddleOCRVL()
        output = pipeline.predict(img_base64)
        for res in output:
            res.print()
            res.save_to_json(save_path="output")
            res.save_to_markdown(save_path="output")

        try:
            # 尝试直接调用API (这里暂时注释掉，因为api_url未定义)
            # response = requests.post(
            #     api_url + "/api/predict",
            #     json={"image": img_base64}
            # )
            # response.raise_for_status()
            # result = response.json()
            # return result.get("text", "")
            raise Exception("API调用暂时未实现")
        except:
            # 如果直接API调用失败，尝试使用备用方案（如使用paddleocr本地库）
            # 这里使用pytesseract作为备用方案
            return pytesseract.image_to_string(image, lang='chi_sim+eng')
    except Exception as e:
        print(f"OCR OlmOCR错误: {e}")
        # 出错时使用pytesseract作为备用
        return pytesseract.image_to_string(image, lang='chi_sim+eng')

def extract_file_content(uploaded_file): # 从上传的文件中提取内容
    if uploaded_file is None:
        return None
    try:
        file_extension = uploaded_file.name.split('.')[-1].lower()
        if file_extension == 'pdf':
            with pdfplumber.open(uploaded_file) as pdf:
                text = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text is None:
                        # 处理扫描PDF
                        image = page.to_image()
                        page_text = ocr_olmocr(image.original)
                    text += page_text
            return text
        elif file_extension in ['png', 'jpeg', 'jpg']:
            image = Image.open(uploaded_file)
            text = ocr_olmocr(image)
            return text
        elif file_extension == 'docx':
            import docx
            doc = docx.Document(uploaded_file)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text
        elif file_extension == 'doc':
            # 处理.doc文件
            import win32com.client as win32
            import tempfile
            import os
            
            # 保存临时文件
            with tempfile.NamedTemporaryFile(delete=False, suffix='.doc') as tmp:
                tmp.write(uploaded_file.read())
                tmp_path = tmp.name
            
            try:
                # 使用Word应用程序打开文件
                word = win32.Dispatch('Word.Application')
                word.Visible = False
                doc = word.Documents.Open(tmp_path)
                text = doc.Content.Text
                doc.Close()
                word.Quit()
                return text
            except:
                # 如果无法使用win32com，尝试转换为docx
                try:
                    import subprocess
                    import os
                    
                    # 使用LibreOffice转换.doc为.docx
                    output_path = tmp_path + 'x'
                    subprocess.run([
                        'soffice', '--headless', '--convert-to', 'docx', 
                        tmp_path, '--outdir', os.path.dirname(tmp_path)
                    ], check=True)
                    
                    # 读取转换后的docx文件
                    import docx
                    doc = docx.Document(output_path)
                    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                    
                    # 删除临时文件
                    os.remove(output_path)
                    return text
                except:
                    print("无法读取.doc文件，请尝试转换为.docx格式后再上传")
                    return None
            finally:
                # 删除临时文件
                os.remove(tmp_path)
        else:
            return uploaded_file.read().decode("utf-8")
    except Exception as e:
        print(f"Error reading file: {e}")
        return None

def get_resume_content(uploaded_file, resume_text):   # 获取简历内容
    """
    获取简历内容
    """
    if uploaded_file is not None:
        return extract_file_content(uploaded_file)
    elif resume_text:
        return resume_text
    else:
        return None

def analyze_resume_with_AI(resume_text: str, jd_text: str, target_position: str, api_key: str) -> str: # 分析简历内容
    """
    分析简历内容
    """
    if not resume_text:
        return "请输入简历内容"
    if not jd_text:
        return "请输入JD内容"
    if not target_position:
        return "请输入目标岗位"
    if not api_key:
        return "请输入你的OpenAI API Key"

    client = OpenAI(
        api_key=api_key,
        base_url="https://dashscope.aliyuncs.com/compatible-mode/v1",
    )
    prompt = f"请基于以下候选人简历文本和目标岗位JD文本，生成一份结构清晰、核心信息用表格呈现的岗位匹配度分析报告，严格遵循以下模块和格式要求：\n\n### 输入材料\n1. 候选人简历文本：{resume_text}\n2. 目标岗位 JD 文本：{jd_text}\n\n### 输出格式要求\n---\n## 一、评分与分析理由板块\n1. **整体评分**：给出0-10分的综合得分\n2. **综合评价**：150字左右的总述，突出匹配亮点与核心短板\n3. **维度拆解分析**：必须用表格呈现，表格列固定为「维度|评分 (/10)|分析理由」，维度包含：\n   - 岗位匹配度\n   - 工作经验相关性\n   - 技能掌握程度\n   - 教育背景契合度\n   - 软技能与岗位适配性\n4. **主要差距总结**：用项目符号列出3-5条最核心的不匹配点\n\n---\n## 二、对照岗位 JD 逐条修改简历板块\n必须用表格呈现，表格列固定为「简历现有内容|岗位 JD 要求|差异分析|修改建议」，需将简历中所有与 JD 相关的条目逐一对应分析，并给出可直接替换的改写话术。\n\n---\n## 三、面试可能问的问题板块\n列出8-10个高针对性问题，每个问题后用「⚠️」标注考察点，例如：\n1. 你在实习中提到「构建多维度测评体系」，能否详细说明你是如何定义「准确性」和「逻辑性」的？⚠️ 考察数据质量把控能力和标准化思维\n\n---\n## 四、职业发展路径板块\n分「短期 (1-3年)」「中期 (3-5年)」「长期 (5年以上)」三个阶段，每个阶段包含：\n- 目标职位\n- 核心任务 / 能力升级重点\n- 行动建议（用「✅」标注具体动作）\n\n---\n## 五、结语建议板块\n给候选人的投递/面试策略总结，3-4条可落地的行动建议。\n\n---\n### 格式约束\n- 所有对比类、评分类内容必须用表格呈现，禁止纯文本堆砌\n- 每个板块用「---」分隔，标题用「#」「##」分级，保持视觉清晰\n- 语言需专业、简洁，避免冗余表述\n\n要求分析全面、具体，避免模板化回复，完全基于提供的JD和简历内容"

    response = client.chat.completions.create(
        model="qwen-plus",
        messages=[
            {"role": "system", "content": "你是一个专业的简历分析助手，根据简历内容给出专业的评分和分析建议"},
            {"role": "user", "content": prompt}
        ],
        max_tokens=1500
    )
    return response.choices[0].message.content

@app.route('/', methods=['GET', 'POST'])
def index():
    """
    主页面路由
    """
    if request.method == 'POST':
        # 获取表单数据
        api_key = request.form.get('api_key')
        target_position = request.form.get('target_position')
        jd_input_method = request.form.get('jd_input_method')
        resume_input_method = request.form.get('resume_input_method')
        
        # 获取JD内容
        jd_text = None
        if jd_input_method == 'paste':
            jd_text = request.form.get('jd_text')
        else:
            jd_file = request.files.get('jd_file')
            if jd_file:
                jd_text = extract_file_content(jd_file)
        
        # 获取简历内容
        resume_text = None
        if resume_input_method == 'paste':
            resume_text = request.form.get('resume_text')
        else:
            resume_file = request.files.get('resume_file')
            if resume_file:
                resume_text = extract_file_content(resume_file)
        
        # 验证输入
        if not resume_text:
            flash('请输入简历内容或上传简历文件', 'warning')
        elif not jd_text:
            flash('请输入JD内容或上传JD文件', 'warning')
        elif not target_position:
            flash('请输入目标岗位', 'warning')
        elif not api_key:
            flash('请输入你的OpenAI API Key', 'warning')
        else:
            try:
                # 分析简历
                analyze_result = analyze_resume_with_AI(resume_text, jd_text, target_position, api_key)
                session['analyze_result'] = analyze_result
                flash('分析完成', 'success')
            except Exception as e:
                flash(f'分析过程中出现错误: {str(e)}', 'danger')
    
    # 获取分析结果
    analyze_result = session.get('analyze_result', None)
    
    # HTML模板
    template = '''
    <!DOCTYPE html>
    <html lang="zh-CN">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>AIPM 简历分析智能助手</title>
        <link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css">
        <style>
            body {
                background-color: #f8f9fa;
                padding: 20px;
            }
            .container {
                max-width: 800px;
                background-color: white;
                border-radius: 10px;
                padding: 30px;
                box-shadow: 0 0 10px rgba(0,0,0,0.1);
                margin-top: 30px;
            }
            h1 {
                color: #343a40;
                margin-bottom: 30px;
                text-align: center;
            }
            .form-group {
                margin-bottom: 20px;
            }
            .btn-primary {
                width: 100%;
                padding: 10px;
                font-size: 16px;
            }
            .result {
                margin-top: 30px;
                padding: 20px;
                background-color: #f8f9fa;
                border-radius: 5px;
                white-space: pre-wrap;
            }
            .alert {
                margin-top: 20px;
            }
        </style>
    </head>
    <body>
        <div class="container">
            <h1>AIPM 简历分析智能助手</h1>
            <p class="text-center text-muted">这是一个基于AI的简历分析工具，帮助你分析简历并给出专业建议。</p>
            
            {% with messages = get_flashed_messages(with_categories=true) %}
                {% if messages %}
                    {% for category, message in messages %}
                        <div class="alert alert-{{ category }}">
                            {{ message }}
                        </div>
                    {% endfor %}
                {% endif %}
            {% endwith %}
            
            <form method="POST" enctype="multipart/form-data">
                <!-- API Key输入 -->
                <div class="form-group">
                    <label for="api_key">请输入你的API Key</label>
                    <input type="password" class="form-control" id="api_key" name="api_key" placeholder="请输入API Key">
                </div>
                
                <!-- 目标岗位输入 -->
                <div class="form-group">
                    <label for="target_position">请输入目标岗位</label>
                    <input type="text" class="form-control" id="target_position" name="target_position" placeholder="例如：AI产品经理">
                </div>
                
                <!-- JD输入部分 -->
                <div class="form-group">
                    <label>请选择JD输入方式</label>
                    <div class="form-check">
                        <input class="form-check-input" type="radio" name="jd_input_method" id="jd_paste" value="paste" checked>
                        <label class="form-check-label" for="jd_paste">粘贴文本</label>
                    </div>
                    <div class="form-check">
                        <input class="form-check-input" type="radio" name="jd_input_method" id="jd_upload" value="upload">
                        <label class="form-check-label" for="jd_upload">上传文件</label>
                    </div>
                </div>
                
                <div class="form-group" id="jd_paste_section">
                    <label for="jd_text">请输入JD内容</label>
                    <textarea class="form-control" id="jd_text" name="jd_text" rows="10" placeholder="请粘贴JD内容"></textarea>
                </div>
                
                <div class="form-group" id="jd_upload_section" style="display: none;">
                    <label for="jd_file">请上传JD文件</label>
                    <input type="file" class="form-control" id="jd_file" name="jd_file" accept=".txt,.md,.pdf,.png,.jpeg,.jpg,.docx,.doc">
                </div>
                
                <!-- 简历输入部分 -->
                <div class="form-group">
                    <label>请选择简历输入方式</label>
                    <div class="form-check">
                        <input class="form-check-input" type="radio" name="resume_input_method" id="resume_paste" value="paste" checked>
                        <label class="form-check-label" for="resume_paste">粘贴文本</label>
                    </div>
                    <div class="form-check">
                        <input class="form-check-input" type="radio" name="resume_input_method" id="resume_upload" value="upload">
                        <label class="form-check-label" for="resume_upload">上传文件</label>
                    </div>
                </div>
                
                <div class="form-group" id="resume_paste_section">
                    <label for="resume_text">请输入简历内容</label>
                    <textarea class="form-control" id="resume_text" name="resume_text" rows="10" placeholder="请粘贴简历内容"></textarea>
                </div>
                
                <div class="form-group" id="resume_upload_section" style="display: none;">
                    <label for="resume_file">请上传你的简历文件</label>
                    <input type="file" class="form-control" id="resume_file" name="resume_file" accept=".txt,.md,.pdf,.png,.jpeg,.jpg,.docx,.doc">
                </div>
                
                <!-- 提交按钮 -->
                <button type="submit" class="btn btn-primary">开始分析</button>
            </form>
            
            <!-- 分析结果 -->
            {% if analyze_result %}
                <div class="result">
                    <h3>分析结果</h3>
                    <pre>{{ analyze_result }}</pre>
                </div>
            {% endif %}
        </div>
        
        <script>
            // JD输入方式切换
            document.querySelectorAll('input[name="jd_input_method"]').forEach(radio => {
                radio.addEventListener('change', function() {
                    if (this.value === 'paste') {
                        document.getElementById('jd_paste_section').style.display = 'block';
                        document.getElementById('jd_upload_section').style.display = 'none';
                    } else {
                        document.getElementById('jd_paste_section').style.display = 'none';
                        document.getElementById('jd_upload_section').style.display = 'block';
                    }
                });
            });
            
            // 简历输入方式切换
            document.querySelectorAll('input[name="resume_input_method"]').forEach(radio => {
                radio.addEventListener('change', function() {
                    if (this.value === 'paste') {
                        document.getElementById('resume_paste_section').style.display = 'block';
                        document.getElementById('resume_upload_section').style.display = 'none';
                    } else {
                        document.getElementById('resume_paste_section').style.display = 'none';
                        document.getElementById('resume_upload_section').style.display = 'block';
                    }
                });
            });
        </script>
        <script src="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/js/bootstrap.bundle.min.js"></script>
    </body>
    </html>
    '''
    
    return render_template_string(template, analyze_result=session.get('analyze_result'))

# 运行Flask应用
if __name__ == '__main__':
    app.run(debug=True, port=5000)
